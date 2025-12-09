import os
import json
import zipfile
import uuid
import hashlib
from typing import Optional, List, Dict, Any
from fastapi import FastAPI, UploadFile, File, Form, Depends, HTTPException, Header
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from docx.shared import RGBColor

# --------------------------------------------------------
# App init
# --------------------------------------------------------
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
    allow_credentials=True
)

# --------------------------------------------------------
# Paths
# --------------------------------------------------------
DATA_DIR = "data"
USERS_FILE = f"{DATA_DIR}/users.json"

os.makedirs(DATA_DIR, exist_ok=True)

# --------------------------------------------------------
# Helper functions
# --------------------------------------------------------

def load_json(path: str) -> Dict[str, Any]:
    if not os.path.exists(path):
        return {}
    try:
        with open(path, "r") as f:
            return json.load(f)
    except:
        return {}

def save_json(path: str, data: Dict[str, Any]):
    with open(path, "w") as f:
        json.dump(data, f, indent=2)

def make_preview(text: str, n: int = 120) -> str:
    text = text.replace("\n", " ")
    text = " ".join(text.split())
    return text[:n] + ("..." if len(text) > n else "")

def hash_password(password: str, salt: str) -> str:
    return hashlib.sha256((salt + password).encode("utf-8")).hexdigest()


def get_user_paths(username: str) -> Dict[str, str]:
    base = os.path.join(DATA_DIR, username)
    os.makedirs(base, exist_ok=True)
    return {
        "docs": os.path.join(base, "documents.json"),
        "anns": os.path.join(base, "annotations.json"),
        "labels": os.path.join(base, "labels.json")
    }


class AuthRequest(BaseModel):
    username: str
    password: str


sessions: Dict[str, str] = {}


def get_current_user(authorization: Optional[str] = Header(None)) -> str:
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing token")
    token = authorization.split(" ", 1)[1].strip()
    username = sessions.get(token)
    if not username:
        raise HTTPException(status_code=401, detail="Invalid token")
    return username

# --------------------------------------------------------
# Auth endpoints
# --------------------------------------------------------

@app.post("/register")
def register_user(payload: AuthRequest):
    username = payload.username.strip()
    password = payload.password
    if not username or not password:
        return JSONResponse({"error": "Username and password required"}, status_code=400)

    users = load_json(USERS_FILE)
    if username in users:
        return JSONResponse({"error": "Username already exists"}, status_code=400)

    salt = uuid.uuid4().hex
    users[username] = {
        "salt": salt,
        "password": hash_password(password, salt)
    }
    save_json(USERS_FILE, users)
    return {"status": "registered"}


@app.post("/login")
def login_user(payload: AuthRequest):
    username = payload.username.strip()
    password = payload.password
    users = load_json(USERS_FILE)
    if username not in users:
        return JSONResponse({"error": "Invalid credentials"}, status_code=401)

    info = users[username]
    expected = info.get("password")
    salt = info.get("salt", "")
    if not expected or hash_password(password, salt) != expected:
        return JSONResponse({"error": "Invalid credentials"}, status_code=401)

    token = uuid.uuid4().hex
    sessions[token] = username
    return {"status": "ok", "token": token, "username": username}

# --------------------------------------------------------
# Upload single file
# --------------------------------------------------------

@app.post("/upload")
async def upload(file: UploadFile = File(...), user: str = Depends(get_current_user)):
    content = await file.read()
    try:
        text = content.decode("utf-8", errors="ignore")
    except:
        text = ""

    paths = get_user_paths(user)
    docs = load_json(paths["docs"])

    doc_id = file.filename
    docs[doc_id] = {
        "doc_id": doc_id,
        "filename": file.filename,
        "text": text,
        "preview": make_preview(text)
    }

    save_json(paths["docs"], docs)
    return {"status": "uploaded", "doc_id": doc_id}

# --------------------------------------------------------
# Upload ZIP
# --------------------------------------------------------

@app.post("/upload-zip")
async def upload_zip(file: UploadFile = File(...), user: str = Depends(get_current_user)):
    paths = get_user_paths(user)
    docs = load_json(paths["docs"])

    data = await file.read()
    path = "temp.zip"
    with open(path, "wb") as f:
        f.write(data)

    with zipfile.ZipFile(path, "r") as z:
        for name in z.namelist():
            if name.endswith(".txt"):
                txt = z.read(name).decode("utf-8", errors="ignore")
                docs[name] = {
                    "doc_id": name,
                    "filename": name,
                    "text": txt,
                    "preview": make_preview(txt)
                }

    save_json(paths["docs"], docs)
    os.remove(path)
    return {"status": "uploaded-zip"}

# --------------------------------------------------------
# Upload folder (multiple files)
# --------------------------------------------------------

@app.post("/upload-folder")
async def upload_folder(
    files: List[UploadFile] = File(...),
    user: str = Depends(get_current_user)
):
    paths = get_user_paths(user)
    docs = load_json(paths["docs"])

    for f in files:
        content = await f.read()
        txt = content.decode("utf-8", errors="ignore")
        docs[f.filename] = {
            "doc_id": f.filename,
            "filename": f.filename,
            "text": txt,
            "preview": make_preview(txt)
        }

    save_json(paths["docs"], docs)
    return {"status": "folder-uploaded"}

# --------------------------------------------------------
# List documents
# --------------------------------------------------------

@app.get("/documents")
def list_docs(user: str = Depends(get_current_user)):
    paths = get_user_paths(user)
    docs = load_json(paths["docs"])
    return list(docs.values())   # IMPORTANT: return list of objects

# --------------------------------------------------------
# Get document text
# --------------------------------------------------------

@app.get("/document/{doc_id}")
def get_doc(doc_id: str, user: str = Depends(get_current_user)):
    paths = get_user_paths(user)
    docs = load_json(paths["docs"])
    if doc_id not in docs:
        return JSONResponse({"error": "not found"}, status_code=404)

    return {
        "doc_id": doc_id,
        "text": docs[doc_id]["text"]
    }

# --------------------------------------------------------
# Save annotation (with rank)
# --------------------------------------------------------

@app.post("/save-annotation")
async def save_annot(
    doc_id: str = Form(...),
    start: int = Form(...),
    end: int = Form(...),
    text: str = Form(...),
    label: str = Form(...),
    rank: Optional[str] = Form(None),
    user: str = Depends(get_current_user)
):
    paths = get_user_paths(user)
    anns = load_json(paths["anns"])
    if doc_id not in anns:
        anns[doc_id] = []

    def overlaps(a_start, a_end, b_start, b_end):
        return a_start < b_end and b_start < a_end

    current = anns[doc_id]
    filtered = [a for a in current if not overlaps(a["start"], a["end"], start, end)]
    filtered.append({
        "start": start,
        "end": end,
        "text": text,
        "label": label,
        "rank": rank
    })

    filtered.sort(key=lambda x: x["start"])
    anns[doc_id] = filtered

    save_json(paths["anns"], anns)
    return {"status": "saved"}

# --------------------------------------------------------
# Get annotations for doc
# --------------------------------------------------------

@app.get("/annotations/{doc_id}")
def get_annots(doc_id: str, user: str = Depends(get_current_user)):
    paths = get_user_paths(user)
    data = load_json(paths["anns"])
    return data.get(doc_id, [])


@app.delete("/annotations/{doc_id}/{index}")
def delete_annotation(doc_id: str, index: int, user: str = Depends(get_current_user)):
    paths = get_user_paths(user)
    anns = load_json(paths["anns"])
    if doc_id not in anns or index < 0 or index >= len(anns[doc_id]):
        return JSONResponse({"error": "annotation not found"}, status_code=404)

    anns[doc_id].pop(index)
    save_json(paths["anns"], anns)
    return {"status": "annotation deleted"}

# --------------------------------------------------------
# Label manager
# --------------------------------------------------------

@app.post("/labels")
async def save_label(
    name: str = Form(...),
    color: str = Form(...),
    user: str = Depends(get_current_user)
):
    paths = get_user_paths(user)
    labels = load_json(paths["labels"])
    labels[name] = color
    save_json(paths["labels"], labels)
    return {"status": "label saved"}


@app.delete("/labels/{label_name}")
async def delete_label(label_name: str, user: str = Depends(get_current_user)):
    paths = get_user_paths(user)
    labels = load_json(paths["labels"])
    if label_name not in labels:
        return JSONResponse({"error": "label not found"}, status_code=404)

    labels.pop(label_name, None)
    save_json(paths["labels"], labels)
    return {"status": "label deleted"}

@app.get("/labels")
def get_label(user: str = Depends(get_current_user)):
    paths = get_user_paths(user)
    return load_json(paths["labels"])

# --------------------------------------------------------
# Export JSON
# --------------------------------------------------------

@app.get("/export-json/{doc_id}")
def export_json_file(doc_id: str, user: str = Depends(get_current_user)):
    paths = get_user_paths(user)
    anns = load_json(paths["anns"]).get(doc_id, [])
    output = f"{doc_id}_annotations.json"
    with open(output, "w") as f:
        json.dump(anns, f, indent=2)
    return FileResponse(output, filename=output)

# --------------------------------------------------------
# Export Word
# --------------------------------------------------------

@app.get("/export-word/{doc_id}")
def export_word(doc_id: str, user: str = Depends(get_current_user)):
    paths = get_user_paths(user)
    docs = load_json(paths["docs"])
    anns = load_json(paths["anns"])

    if doc_id not in docs:
        return JSONResponse({"error": "Not found"}, status_code=404)

    document = Document()
    document.add_heading(f"Annotations for {doc_id}", level=1)

    for a in anns.get(doc_id, []):
        p = document.add_paragraph()
        run = p.add_run(f"[{a['label']}] {a['text']} (Rank={a.get('rank','')})")
        run.font.color.rgb = RGBColor(200, 0, 0)

    filename = f"{doc_id}_annotations.docx"
    document.save(filename)

    return FileResponse(filename, filename)
